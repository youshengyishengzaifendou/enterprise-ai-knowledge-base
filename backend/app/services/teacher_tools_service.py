from pathlib import Path
import re
import uuid

from sqlalchemy.orm import Session

from app.models import KnowledgeChunk, KnowledgeDocument, User
from app.schemas.agent_tools import Citation
from app.services.knowledge_service import search_knowledge


PROJECT_ROOT = Path(__file__).resolve().parents[3]
TEACHER_EXPORT_ROOT = PROJECT_ROOT / "uploads" / "teacher_exports"


def search_teacher_materials(
    db: Session,
    *,
    query: str,
    user: User,
    limit: int = 5,
) -> dict:
    matches = search_knowledge(db, query=query, user=user, limit=_clamp(limit, 1, 10))
    materials = [_material_from_match(row) for row in matches]
    return {
        "materials": materials,
        "citations": _citations_from_matches(matches),
    }


def generate_teacher_questions(
    db: Session,
    *,
    payload: dict,
    user: User,
) -> dict:
    topic = _clean_text(payload.get("topic")) or _clean_text(payload.get("query")) or "知识点"
    question_type = _clean_text(payload.get("question_type")) or "简答题"
    difficulty = _clean_text(payload.get("difficulty")) or "中等"
    count = _clamp(int(payload.get("count") or 3), 1, 20)
    material_id = str(payload.get("material_id") or "").strip()
    matches = _matches_for_teacher_generation(db, topic=topic, user=user, material_id=material_id, limit=max(count, 3))
    context_items = [_context_item_from_match(row) for row in matches]
    questions = [
        _build_question(
            topic=topic,
            question_type=question_type,
            difficulty=difficulty,
            index=index,
            context=context_items[index % len(context_items)] if context_items else None,
        )
        for index in range(count)
    ]
    return {
        "topic": topic,
        "question_type": question_type,
        "difficulty": difficulty,
        "questions": questions,
        "materials": context_items,
        "citations": _citations_from_matches(matches),
    }


def generate_teacher_paper(
    db: Session,
    *,
    payload: dict,
    user: User,
) -> dict:
    topic = _clean_text(payload.get("topic")) or "知识点"
    title = _clean_text(payload.get("title")) or f"{topic}测试卷"
    duration_minutes = _clamp(int(payload.get("duration_minutes") or 45), 5, 240)
    include_answers = bool(payload.get("include_answers", True))
    include_analysis = bool(payload.get("include_analysis", True))
    question_counts = _question_counts(payload.get("question_counts"))
    matches = _matches_for_teacher_generation(db, topic=topic, user=user, limit=max(sum(question_counts.values()), 5))
    context_items = [_context_item_from_match(row) for row in matches]

    sections = []
    question_index = 0
    for question_type, count in question_counts.items():
        questions = []
        for _ in range(count):
            context = context_items[question_index % len(context_items)] if context_items else None
            questions.append(
                _build_question(
                    topic=topic,
                    question_type=question_type,
                    difficulty=_clean_text(payload.get("difficulty")) or "中等",
                    index=question_index,
                    context=context,
                )
            )
            question_index += 1
        sections.append({"question_type": question_type, "count": count, "questions": questions})

    paper = {
        "title": title,
        "topic": topic,
        "paper_type": _clean_text(payload.get("paper_type")) or "quiz",
        "duration_minutes": duration_minutes,
        "include_answers": include_answers,
        "include_analysis": include_analysis,
        "sections": sections,
        "materials": context_items,
    }
    download = _write_teacher_docx(title=title, paragraphs=_paper_paragraphs(paper))
    return {"paper": paper, "download": download, "citations": _citations_from_matches(matches)}


def export_teacher_knowledge(
    db: Session,
    *,
    payload: dict,
    user: User,
) -> dict:
    topic = _clean_text(payload.get("topic")) or "知识点"
    title = _clean_text(payload.get("title")) or f"{topic}知识点整理"
    matches = _matches_for_teacher_generation(db, topic=topic, user=user, limit=5)
    context_items = [_context_item_from_match(row) for row in matches]
    edited_points = payload.get("edited_points")
    if isinstance(edited_points, list) and edited_points:
        points = [_clean_text(item) for item in edited_points if _clean_text(item)]
    else:
        points = [_point_from_context(item["snippet"]) for item in context_items]
    if not points:
        points = [f"围绕{topic}补充核心概念、常见题型和易错点。"]

    handout = {
        "title": title,
        "topic": topic,
        "stage": _clean_text(payload.get("stage")) or "all",
        "subject": _clean_text(payload.get("subject")) or "all",
        "points": points,
        "materials": context_items,
    }
    download = _write_teacher_docx(title=title, paragraphs=_handout_paragraphs(handout))
    return {"handout": handout, "download": download, "citations": _citations_from_matches(matches)}


def _matches_for_teacher_generation(
    db: Session,
    *,
    topic: str,
    user: User,
    limit: int = 5,
    material_id: str = "",
) -> list[dict]:
    if material_id:
        document = db.get(KnowledgeDocument, material_id)
        if document is not None:
            chunks = db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == document.id).order_by(KnowledgeChunk.chunk_index.asc()).all()
            return [
                {
                    "document": document,
                    "chunk": chunk,
                    "score": 1.0,
                    "keyword_score": 1.0,
                    "vector_score": 0.0,
                    "retrieval_method": "material_id",
                }
                for chunk in chunks[:limit]
            ]
    return search_knowledge(db, query=topic, user=user, limit=_clamp(limit, 1, 20))


def _material_from_match(row: dict) -> dict:
    document: KnowledgeDocument = row["document"]
    chunk: KnowledgeChunk = row["chunk"]
    return {
        "document_id": document.id,
        "document_title": document.title,
        "chunk_id": chunk.id,
        "chunk_index": chunk.chunk_index,
        "snippet": chunk.content_text,
        "score": row["score"],
        "source_file_name": document.source_file_name,
        "source_url": document.source_url,
        "document_version": document.current_version,
    }


def _context_item_from_match(row: dict) -> dict:
    material = _material_from_match(row)
    return {
        "document_id": material["document_id"],
        "document_title": material["document_title"],
        "chunk_id": material["chunk_id"],
        "snippet": material["snippet"],
    }


def _build_question(
    *,
    topic: str,
    question_type: str,
    difficulty: str,
    index: int,
    context: dict | None,
) -> dict:
    snippet = context["snippet"] if context else f"{topic}的核心概念和典型应用。"
    source_title = context["document_title"] if context else "未匹配资料"
    key_point = _point_from_context(snippet)
    stem = f"{index + 1}. 【{difficulty}】围绕{topic}，根据“{key_point}”设计一道{question_type}。"
    answer = _answer_for_question(question_type=question_type, topic=topic, key_point=key_point)
    return {
        "id": f"q-{index + 1}",
        "topic": topic,
        "question_type": question_type,
        "difficulty": difficulty,
        "stem": stem,
        "answer": answer,
        "analysis": f"依据资料《{source_title}》中的内容生成，重点检查学生对“{key_point}”的理解。",
        "source_document_title": source_title,
    }


def _answer_for_question(*, question_type: str, topic: str, key_point: str) -> str:
    if "选择" in question_type:
        return f"参考答案：选择能体现“{key_point}”的选项。"
    if "填空" in question_type:
        return f"参考答案：{key_point}"
    return f"参考答案要点：围绕{topic}说明{key_point}，并结合题目条件作答。"


def _point_from_context(text: str) -> str:
    normalized = " ".join(str(text).split())
    first = re.split(r"[。！？.!?；;]", normalized, maxsplit=1)[0].strip()
    if not first:
        return "核心知识点"
    return first[:80]


def _question_counts(value: object) -> dict[str, int]:
    if isinstance(value, dict) and value:
        counts = {}
        for key, count in value.items():
            label = _clean_text(key) or "简答题"
            try:
                counts[label] = _clamp(int(count), 0, 50)
            except (TypeError, ValueError):
                continue
        counts = {key: count for key, count in counts.items() if count > 0}
        if counts:
            return counts
    return {"选择题": 5, "填空题": 3, "简答题": 2}


def _write_teacher_docx(*, title: str, paragraphs: list[str]) -> dict:
    TEACHER_EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^A-Za-z0-9\u4e00-\u9fff_.-]+", "-", title).strip("-") or "teacher-export"
    filename = f"{safe_title}-{uuid.uuid4()}.docx"
    output_path = TEACHER_EXPORT_ROOT / filename
    try:
        from docx import Document

        document = Document()
        document.add_heading(title, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(output_path)
    except Exception:
        output_path.write_text("\n\n".join([title, *paragraphs]), encoding="utf-8")
    return {
        "format": "docx",
        "file_name": filename,
        "path": str(output_path),
        "url": f"/uploads/teacher_exports/{filename}",
    }


def _paper_paragraphs(paper: dict) -> list[str]:
    paragraphs = [f"考试时长：{paper['duration_minutes']}分钟", f"知识点：{paper['topic']}"]
    for section in paper["sections"]:
        paragraphs.append(f"{section['question_type']}（共{section['count']}题）")
        for question in section["questions"]:
            paragraphs.append(question["stem"])
            if paper["include_answers"]:
                paragraphs.append(question["answer"])
            if paper["include_analysis"]:
                paragraphs.append(question["analysis"])
    return paragraphs


def _handout_paragraphs(handout: dict) -> list[str]:
    paragraphs = [f"知识点：{handout['topic']}", f"学段：{handout['stage']}，学科：{handout['subject']}"]
    paragraphs.extend(f"- {point}" for point in handout["points"])
    return paragraphs


def _citations_from_matches(matches: list[dict]) -> list[Citation]:
    return [
        Citation(
            type="knowledge_chunk",
            id=row["chunk"].id,
            title=f"{row['document'].title}#片段{row['chunk'].chunk_index + 1}",
            updated_at=row["document"].updated_at,
        )
        for row in matches
    ]


def _clean_text(value: object) -> str:
    return str(value or "").strip()


def _clamp(value: int, minimum: int, maximum: int) -> int:
    return max(minimum, min(value, maximum))
