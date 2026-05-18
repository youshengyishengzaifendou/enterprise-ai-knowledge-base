import base64
from io import BytesIO
from pathlib import Path

from sqlalchemy import create_engine
from sqlalchemy.orm import Session
from openpyxl import Workbook
from docx import Document

from app.db.base import Base
from app.models import (
    AuditLog,
    Customer,
    KnowledgeChunk,
    KnowledgeDocument,
    KnowledgeDocumentGroup,
    KnowledgeDocumentGroupMembership,
    KnowledgeDocumentPermission,
    KnowledgeDocumentVersion,
    KnowledgeConflictReport,
    KnowledgeProjectPermission,
    KnowledgeSourceFile,
    KnowledgeSourceFilePermission,
    Project,
    User,
    UserChannelBinding,
)
from app.schemas.agent_tools import Actor
from app.services.identity_service import resolve_actor
from app.services.knowledge_service import (
    add_document_to_group,
    answer_with_knowledge,
    batch_grant_documents,
    can_access_source_file,
    create_document_group,
    find_source_files,
    grant_project_permission,
    grant_source_file_permission,
    ingest_document,
    backfill_source_files,
    publish_document,
    reject_document,
    rebuild_document_index,
    search_knowledge,
)
from app.services.permission_operations_service import disable_user, grant_document_permission
from app.services.support_operations_service import (
    get_support_operations_dashboard,
    import_knowledge_file,
    import_faq_text,
    list_unanswered_questions,
    update_unanswered_question_status,
)
from app.services.teacher_tools_service import export_teacher_knowledge, generate_teacher_paper, generate_teacher_questions, search_teacher_materials


def make_session() -> Session:
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    return Session(engine)


def test_existing_document_defaults_to_published_review_state():
    db = make_session()
    document = KnowledgeDocument(
        title="历史知识",
        source_type="manual",
        content_text="历史知识默认仍可检索。",
        created_by="agent-1",
    )
    db.add(document)
    db.commit()

    saved = db.get(KnowledgeDocument, document.id)

    assert saved.review_status == "approved"
    assert saved.publication_status == "published"
    assert saved.current_version == 1


def test_source_file_tracks_parse_status_and_linked_document():
    db = make_session()
    source_file = KnowledgeSourceFile(
        file_name="原件.pdf",
        file_path="/tmp/source.pdf",
        uploaded_by="agent-1",
    )
    db.add(source_file)
    db.commit()

    saved = db.get(KnowledgeSourceFile, source_file.id)

    assert saved.parse_status == "pending"
    assert saved.linked_document_id is None


def test_pending_review_document_is_not_searchable_until_published():
    db = make_session()
    document = ingest_document(
        db,
        payload={
            "title": "待审核退款规则",
            "content_text": "待审核内容不应该被客服检索。",
            "review_status": "pending_review",
            "publication_status": "draft",
            "index_status": "pending_review",
        },
        user_id="agent-1",
    )

    assert search_knowledge(db, query="待审核内容", user=None) == []

    published = publish_document(db, document.id, reviewer_user_id="admin")
    matches = search_knowledge(db, query="待审核内容", user=None)

    assert published["document"]["review_status"] == "approved"
    assert published["document"]["publication_status"] == "published"
    assert published["version"]["version_number"] == 1
    assert matches


def test_answer_citations_include_document_version():
    db = make_session()
    document = ingest_document(
        db,
        payload={"title": "退货政策版本", "content_text": "七天内可以退货。"},
        user_id="agent-1",
    )
    publish_document(db, document.id, reviewer_user_id="admin")

    result = answer_with_knowledge(db, query="几天内可以退货", user=None)

    assert result["matches"][0]["document_version"] == 1


def test_import_knowledge_file_preserves_source_file_when_parse_fails(tmp_path, monkeypatch):
    from app.services import support_operations_service

    db = make_session()
    monkeypatch.setattr(support_operations_service, "UPLOAD_ROOT", tmp_path / "knowledge_sources")

    result = import_knowledge_file(db, filename="空白.pdf", content=b"not-a-real-pdf", user_id="agent-1")

    document = db.query(KnowledgeDocument).one()
    source_file = db.query(KnowledgeSourceFile).one()
    assert result["ok"] is False
    assert document.parse_status == "parse_failed"
    assert document.publication_status == "draft"
    assert document.review_status == "pending_review"
    assert document.index_status == "parse_failed"
    assert source_file.parse_status == "parse_failed"
    assert source_file.linked_document_id == document.id
    assert list((tmp_path / "knowledge_sources").glob("*.pdf"))


def test_publish_document_records_conflict_report_for_opposing_similar_policy():
    db = make_session()
    ingest_document(
        db,
        payload={"title": "退货政策A", "content_text": "客户七天内可以退货。"},
        user_id="agent-1",
    )
    incoming = ingest_document(
        db,
        payload={
            "title": "退货政策B",
            "content_text": "客户七天内不可以退货。",
            "review_status": "pending_review",
            "publication_status": "draft",
            "index_status": "pending_review",
        },
        user_id="agent-1",
    )

    result = publish_document(db, incoming.id, reviewer_user_id="admin")

    assert result["similar_documents"]
    assert result["conflict_reports"]
    assert db.query(KnowledgeConflictReport).count() == 1


def test_unanswered_question_is_captured_when_answer_has_no_match():
    db = make_session()

    result = answer_with_knowledge(db, query="发票抬头写错了怎么处理", user=None, actor_user_id="agent-1")

    unanswered = list_unanswered_questions(db)
    assert result["answer"].startswith("没有找到可用的知识库内容")
    assert len(unanswered) == 1
    assert unanswered[0]["question"] == "发票抬头写错了怎么处理"
    assert unanswered[0]["status"] == "pending"


def test_dashboard_reports_knowledge_operations_metrics():
    db = make_session()
    user = User(id="agent-1", name="客服A", role="member", status="active")
    document = KnowledgeDocument(
        id="doc-1",
        title="退款政策",
        source_type="manual",
        summary="退款处理规则",
        content_text="订单未发货可以直接退款，已发货需先拒收或退回。",
        created_by="agent-1",
    )
    log_hit = AuditLog(id="audit-hit", user_id="agent-1", channel="openclaw", action="query", tool_name="kb_answer", response_summary="已根据知识库生成回答。")
    log_hit.request_payload = {"input": {"query": "订单未发货怎么退款"}}
    log_miss = AuditLog(id="audit-miss", user_id="agent-1", channel="openclaw", action="query", tool_name="kb_answer", response_summary="没有找到可用的知识库内容。")
    log_miss.request_payload = {"input": {"query": "发票抬头写错了怎么办"}}
    db.add_all([user, document, log_hit, log_miss])
    db.commit()
    answer_with_knowledge(db, query="发票抬头写错了怎么办", user=None, actor_user_id="agent-1")

    dashboard = get_support_operations_dashboard(db)

    assert dashboard["metrics"]["knowledge_documents"] == 1
    assert dashboard["metrics"]["total_queries"] == 2
    assert dashboard["metrics"]["hit_queries"] == 1
    assert dashboard["metrics"]["missed_queries"] == 1
    assert dashboard["metrics"]["unanswered_questions"] == 1
    assert dashboard["popular_questions"][0]["question"] == "订单未发货怎么退款"


def test_dashboard_reports_review_parse_conflict_and_source_file_queues():
    db = make_session()
    pending_document = KnowledgeDocument(
        id="doc-pending",
        title="待审核发票规则",
        source_type="file_upload",
        content_text="发票规则等待管理员审核。",
        created_by="agent-1",
        review_status="pending_review",
        publication_status="draft",
        index_status="pending_review",
    )
    approved_document = KnowledgeDocument(
        id="doc-approved",
        title="已发布发票规则",
        source_type="manual",
        content_text="发票抬头错误可以重开。",
        created_by="agent-1",
    )
    failed_file = KnowledgeSourceFile(
        id="file-failed",
        file_name="无法解析.pdf",
        file_path="/tmp/无法解析.pdf",
        uploaded_by="agent-1",
        parse_status="parse_failed",
        parse_error="failed to parse file",
        linked_document_id="doc-pending",
    )
    parsed_file = KnowledgeSourceFile(
        id="file-parsed",
        file_name="已解析.docx",
        file_path="/tmp/已解析.docx",
        uploaded_by="agent-1",
        parse_status="parsed",
        linked_document_id="doc-approved",
    )
    conflict = KnowledgeConflictReport(
        document_id="doc-pending",
        related_document_id="doc-approved",
        detail="当前文档包含冲突条款。",
    )
    db.add_all([pending_document, approved_document, failed_file, parsed_file, conflict])
    db.commit()

    dashboard = get_support_operations_dashboard(db)

    assert dashboard["metrics"]["pending_review_documents"] == 1
    assert dashboard["metrics"]["parse_failed_files"] == 1
    assert dashboard["metrics"]["open_conflicts"] == 1
    assert dashboard["pending_review_documents"][0]["id"] == "doc-pending"
    assert dashboard["parse_failed_files"][0]["id"] == "file-failed"
    assert dashboard["recent_source_files"][0]["file_name"] in {"无法解析.pdf", "已解析.docx"}
    assert dashboard["conflict_reports"][0]["related_document_title"] == "已发布发票规则"


def test_import_faq_text_creates_support_knowledge_documents():
    db = make_session()

    result = import_faq_text(
        db,
        text="问题,答案\n怎么退款,订单未发货可以直接退款\n物流丢件怎么办,先联系快递核实并补发",
        user_id="agent-1",
        source_type="faq_csv",
    )

    assert result["imported_count"] == 2
    documents = db.query(KnowledgeDocument).order_by(KnowledgeDocument.title.asc()).all()
    assert [document.title for document in documents] == ["怎么退款", "物流丢件怎么办"]
    assert documents[0].source_type == "faq_csv"


def test_ingest_document_preserves_source_file_index():
    db = make_session()

    document = ingest_document(
        db,
        payload={
            "title": "邮件安全升级手册",
            "content_text": "升级前需要备份配置，并在维护窗口执行升级。",
            "source_file_path": "/root/.openclaw/media/inbound/邮件安全升级手册.pdf",
            "source_file_name": "邮件安全升级手册.pdf",
            "source_file_mime_type": "application/pdf",
            "source_file_size": 2048,
            "source_file_storage": "openclaw_media",
        },
        user_id="agent-1",
    )

    saved = db.get(KnowledgeDocument, document.id)
    assert saved is not None
    assert saved.source_file_path == "/root/.openclaw/media/inbound/邮件安全升级手册.pdf"
    assert saved.source_file_name == "邮件安全升级手册.pdf"
    assert saved.source_file_mime_type == "application/pdf"
    assert saved.source_file_size == 2048
    assert saved.source_file_storage == "openclaw_media"


def test_ingest_document_records_source_file_and_source_account():
    db = make_session()

    document = ingest_document(
        db,
        payload={
            "title": "飞书上传原件",
            "content_text": "飞书上传的原文档需要能追溯来源账号。",
            "source_file_path": "/root/.openclaw/media/inbound/feishu.docx",
            "source_file_name": "feishu.docx",
            "source_file_storage": "openclaw_media",
        },
        user_id="agent-1",
        source_channel="feishu",
        source_external_user_id="ou_demo",
    )

    source_file = db.query(KnowledgeSourceFile).one()
    saved = db.get(KnowledgeDocument, document.id)
    assert saved is not None
    assert saved.source_file_id == source_file.id
    assert saved.source_channel == "feishu"
    assert saved.source_external_user_id == "ou_demo"
    assert source_file.file_name == "feishu.docx"
    assert source_file.file_path == "/root/.openclaw/media/inbound/feishu.docx"
    assert source_file.uploaded_by == "agent-1"
    assert source_file.source_channel == "feishu"
    assert source_file.source_external_user_id == "ou_demo"


def test_ingest_document_indexes_chunks_for_hybrid_retrieval():
    db = make_session()

    document = ingest_document(
        db,
        payload={
            "title": "发票处理规则",
            "content_text": "发票抬头错误时，未开票可以修改，已开票需要作废重开。",
        },
        user_id="agent-1",
    )

    chunk = db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == document.id).one()
    assert chunk.embedding_model == "local-hash-v1"
    assert chunk.embedding_json
    assert chunk.index_status == "indexed"
    assert chunk.indexed_at is not None
    assert chunk.token_count > 0
    assert document.parse_status == "parsed"
    assert document.index_status == "indexed"


def test_rebuild_document_index_reindexes_missing_embeddings():
    db = make_session()
    document = ingest_document(
        db,
        payload={
            "title": "物流异常处理",
            "content_text": "客户反馈物流停滞超过48小时，客服需要联系快递核实并同步客户。",
        },
        user_id="agent-1",
    )
    chunk = db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == document.id).one()
    chunk.embedding_json = None
    chunk.index_status = "not_indexed"
    chunk.indexed_at = None
    db.commit()

    result = rebuild_document_index(db, document.id)
    rebuilt = db.get(KnowledgeChunk, chunk.id)

    assert result["indexed_chunks"] == 1
    assert result["index_status"] == "indexed"
    assert rebuilt is not None
    assert rebuilt.embedding_json
    assert rebuilt.index_status == "indexed"
    assert rebuilt.indexed_at is not None


def test_search_knowledge_filters_by_user_project_permissions():
    db = make_session()
    owner = User(id="owner", name="负责人", role="member", status="active")
    outsider = User(id="outsider", name="外部坐席", role="member", status="active")
    customer = Customer(id="customer-1", name="客户A", status="active", owner_user_id="owner")
    project = Project(id="project-1", customer_id="customer-1", name="客户A售后", status="active", owner_user_id="owner")
    db.add_all([owner, outsider, customer, project])
    db.commit()
    ingest_document(
        db,
        payload={
            "title": "客户A退款规则",
            "content_text": "客户A订单未发货可以直接退款。",
            "project_id": "project-1",
            "customer_id": "customer-1",
        },
        user_id="owner",
    )

    owner_matches = search_knowledge(db, query="客户A怎么退款", user=owner)
    outsider_matches = search_knowledge(db, query="客户A怎么退款", user=outsider)

    assert owner_matches
    assert outsider_matches == []


def test_knowledge_admin_can_search_unscoped_documents_created_by_other_users():
    db = make_session()
    admin = User(id="admin-user", name="历史全库账号", role="member", status="active", knowledge_access_policy="all")
    private_user = User(id="private-user", name="新增微信账号", role="member", status="active", knowledge_access_policy="own")
    db.add_all([admin, private_user])
    db.commit()
    ingest_document(
        db,
        payload={
            "title": "全库可见资料",
            "content_text": "历史全库账号可以查询这篇资料。",
        },
        user_id="private-user",
    )

    matches = search_knowledge(db, query="历史全库账号可以查询什么", user=admin)

    assert matches
    assert matches[0]["document"].created_by == "private-user"


def test_new_channel_actor_is_auto_created_with_private_knowledge_policy():
    db = make_session()

    user = resolve_actor(db, Actor(channel="openclaw-weixin", external_user_id="new-wx-user@im.wechat"))

    assert user is not None
    assert user.role == "member"
    assert user.knowledge_access_policy == "own"
    binding = db.query(UserChannelBinding).filter_by(channel="openclaw-weixin", external_user_id="new-wx-user@im.wechat").one()
    assert binding.user_id == user.id


def test_weixin_account_alias_is_used_as_knowledge_user_id():
    db = make_session()

    user = resolve_actor(db, Actor(channel="openclaw-weixin", external_user_id="wx-zhupeiting"))

    assert user is not None
    assert user.id == "wx-zhupeiting"
    assert user.name == "openclaw-weixin:wx-zhupeiting"
    assert user.knowledge_access_policy == "own"
    binding = db.query(UserChannelBinding).filter_by(channel="openclaw-weixin", external_user_id="wx-zhupeiting").one()
    assert binding.user_id == "wx-zhupeiting"


def test_existing_channel_user_id_can_be_resolved_without_binding():
    db = make_session()
    db.add(User(id="wx-wangsuzhen", name="openclaw-weixin:wx-wangsuzhen", role="member", status="active", knowledge_access_policy="all"))
    db.commit()

    user = resolve_actor(db, Actor(channel="openclaw-weixin", external_user_id="wx-wangsuzhen"))

    assert user is not None
    assert user.id == "wx-wangsuzhen"
    assert user.knowledge_access_policy == "all"
    binding = db.query(UserChannelBinding).filter_by(channel="openclaw-weixin", external_user_id="wx-wangsuzhen").one()
    assert binding.user_id == "wx-wangsuzhen"


def test_unknown_legacy_chat_actor_can_use_demo_fallback_binding():
    db = make_session()
    db.add(User(id="user-demo", name="演示用户", role="project_owner", status="active", knowledge_access_policy="all"))
    db.add(UserChannelBinding(user_id="user-demo", channel="openclaw", external_user_id="unknown"))
    db.commit()

    user = resolve_actor(db, Actor(channel="openclaw-weixin", external_user_id="unknown"))

    assert user is not None
    assert user.id == "user-demo"
    assert user.knowledge_access_policy == "all"


def test_configured_legacy_global_actor_is_auto_created_with_global_policy(monkeypatch):
    from app.services import identity_service

    monkeypatch.setattr(identity_service, "LEGACY_GLOBAL_ACTORS", {("feishu", "configured_legacy")})
    db = make_session()

    user = resolve_actor(db, Actor(channel="feishu", external_user_id="configured_legacy"))

    assert user is not None
    assert user.knowledge_access_policy == "all"


def test_private_users_only_search_own_or_granted_unscoped_documents():
    db = make_session()
    owner = User(id="owner-user", name="文档上传人", role="member", status="active", knowledge_access_policy="own")
    outsider = User(id="outsider-user", name="新增账号", role="member", status="active", knowledge_access_policy="own")
    db.add_all([owner, outsider])
    db.commit()
    owner_document = ingest_document(
        db,
        payload={
            "title": "私有售后资料",
            "content_text": "私有资料说明订单未发货可以退款。",
        },
        user_id="owner-user",
    )
    outsider_document = ingest_document(
        db,
        payload={
            "title": "外部账号资料",
            "content_text": "外部账号自己的知识说明发票可重开。",
        },
        user_id="outsider-user",
    )

    owner_matches = search_knowledge(db, query="订单未发货可以退款", user=owner)
    outsider_matches_before_grant = search_knowledge(db, query="订单未发货可以退款", user=outsider)
    db.add(KnowledgeDocumentPermission(document_id=owner_document.id, user_id="outsider-user", access_level="read"))
    db.commit()
    outsider_matches_after_grant = search_knowledge(db, query="订单未发货可以退款", user=outsider)
    outsider_own_matches = search_knowledge(db, query="发票可重开", user=outsider)

    assert [match["document"].id for match in owner_matches] == [owner_document.id]
    assert outsider_matches_before_grant == []
    assert [match["document"].id for match in outsider_matches_after_grant] == [owner_document.id]
    assert outsider_document.id in [match["document"].id for match in outsider_own_matches]


def test_find_source_files_respects_private_document_permissions():
    db = make_session()
    owner = User(id="source-owner", name="文件上传人", role="member", status="active", knowledge_access_policy="own")
    outsider = User(id="source-outsider", name="新增账号", role="member", status="active", knowledge_access_policy="own")
    db.add_all([owner, outsider])
    db.commit()
    document = ingest_document(
        db,
        payload={
            "title": "授权原文档",
            "content_text": "这份资料有原文档。",
            "source_file_path": "/root/.openclaw/media/inbound/private.docx",
            "source_file_name": "private.docx",
            "source_file_storage": "openclaw_media",
        },
        user_id="source-owner",
    )

    assert find_source_files(db, query="private", user=outsider) == []
    db.add(KnowledgeDocumentPermission(document_id=document.id, user_id="source-outsider", access_level="read"))
    db.commit()

    files = find_source_files(db, query="private", user=outsider)

    assert len(files) == 1
    assert files[0]["document_id"] == document.id


def test_source_file_permission_can_grant_original_file_without_document_permission():
    db = make_session()
    owner = User(id="source-owner", name="文件上传人", role="member", status="active", knowledge_access_policy="own")
    viewer = User(id="source-viewer", name="文件查看人", role="member", status="active", knowledge_access_policy="own")
    db.add_all([owner, viewer])
    db.commit()
    document = ingest_document(
        db,
        payload={
            "title": "只授权原件",
            "content_text": "原件可单独授权。",
            "source_file_path": "/root/.openclaw/media/inbound/source-only.pdf",
            "source_file_name": "source-only.pdf",
            "source_file_storage": "openclaw_media",
        },
        user_id="source-owner",
    )
    source_file = db.query(KnowledgeSourceFile).one()

    assert can_access_source_file(db, viewer, source_file) is False
    permission = grant_source_file_permission(db, source_file_id=source_file.id, user_id="source-viewer", access_level="read")
    files = find_source_files(db, query="source-only", user=viewer)

    assert permission.source_file_id == source_file.id
    assert files[0]["source_file_id"] == source_file.id
    assert files[0]["document_id"] == document.id


def test_backfill_source_files_links_existing_documents():
    db = make_session()
    db.add(
        KnowledgeDocument(
            id="legacy-doc",
            title="旧原文档",
            source_type="manual",
            source_file_path="/root/.openclaw/media/inbound/legacy.pdf",
            source_file_name="legacy.pdf",
            source_file_storage="openclaw_media",
            content_text="旧知识已有原文档路径。",
            created_by="legacy-user",
        )
    )
    db.commit()

    result = backfill_source_files(db)
    document = db.get(KnowledgeDocument, "legacy-doc")
    source_file = db.query(KnowledgeSourceFile).one()

    assert result["backfilled_count"] == 1
    assert document.source_file_id == source_file.id
    assert source_file.file_path == "/root/.openclaw/media/inbound/legacy.pdf"


def test_existing_feishu_and_weixin_actors_are_auto_created_with_global_knowledge_policy():
    db = make_session()

    feishu_user = resolve_actor(db, Actor(channel="feishu", external_user_id="ou_8d05034bd270234aff8cdefa87f2a5ba"))
    weixin_user = resolve_actor(db, Actor(channel="openclaw-weixin", external_user_id="o9cq801FsEsSNQ-z6MR_xQz6yb1Q@im.wechat"))

    assert feishu_user is not None
    assert weixin_user is not None
    assert feishu_user.knowledge_access_policy == "all"
    assert weixin_user.knowledge_access_policy == "all"


def test_document_group_membership_records_document_grouping():
    db = make_session()
    group = create_document_group(db, name="售后政策", created_by="admin-user", description="售后相关文档")
    document = ingest_document(
        db,
        payload={"title": "退货政策", "content_text": "七天内可申请退货。"},
        user_id="admin-user",
    )

    membership = add_document_to_group(db, group_id=group.id, document_id=document.id)

    assert db.get(KnowledgeDocumentGroup, group.id).name == "售后政策"
    assert db.get(KnowledgeDocumentGroupMembership, membership.id).document_id == document.id


def test_project_permission_grants_project_scoped_knowledge_access():
    db = make_session()
    owner = User(id="project-owner", name="项目负责人", role="member", status="active", knowledge_access_policy="own")
    viewer = User(id="project-viewer", name="项目查看人", role="member", status="active", knowledge_access_policy="own")
    customer = Customer(id="customer-project", name="项目客户", status="active", owner_user_id="project-owner")
    project = Project(id="project-permission", customer_id="customer-project", name="项目知识", status="active", owner_user_id="project-owner")
    db.add_all([owner, viewer, customer, project])
    db.commit()
    document = ingest_document(
        db,
        payload={
            "title": "项目级知识",
            "content_text": "项目级授权后可以看到这条知识。",
            "project_id": "project-permission",
            "customer_id": "customer-project",
        },
        user_id="project-owner",
    )

    assert search_knowledge(db, query="项目级授权后", user=viewer) == []
    permission = grant_project_permission(db, project_id="project-permission", user_id="project-viewer")
    matches = search_knowledge(db, query="项目级授权后", user=viewer)

    assert permission.project_id == "project-permission"
    assert matches[0]["document"].id == document.id


def test_batch_grant_documents_grants_multiple_documents_to_user():
    db = make_session()
    owner = User(id="batch-owner", name="批量上传人", role="member", status="active", knowledge_access_policy="own")
    viewer = User(id="batch-viewer", name="批量查看人", role="member", status="active", knowledge_access_policy="own")
    db.add_all([owner, viewer])
    db.commit()
    first = ingest_document(db, payload={"title": "批量授权一", "content_text": "批量授权第一篇。"}, user_id="batch-owner")
    second = ingest_document(db, payload={"title": "批量授权二", "content_text": "批量授权第二篇。"}, user_id="batch-owner")

    permissions = batch_grant_documents(db, document_ids=[first.id, second.id], user_id="batch-viewer")
    matches = search_knowledge(db, query="批量授权", user=viewer, limit=10)

    assert {permission.document_id for permission in permissions} == {first.id, second.id}
    assert {match["document"].id for match in matches} >= {first.id, second.id}


def test_permission_grant_writes_audit_log():
    db = make_session()
    owner = User(id="audit-owner", name="审计授权人", role="member", status="active", knowledge_access_policy="all")
    viewer = User(id="audit-viewer", name="审计查看人", role="member", status="active", knowledge_access_policy="own")
    db.add_all([owner, viewer])
    db.commit()
    document = ingest_document(db, payload={"title": "审计知识", "content_text": "授权需要审计。"}, user_id="audit-owner")

    permission = grant_document_permission(db, document_id=document.id, user_id="audit-viewer", actor_user_id="audit-owner")
    audit = db.query(AuditLog).filter(AuditLog.tool_name == "knowledge_permission_grant").one()

    assert permission.document_id == document.id
    assert audit.user_id == "audit-owner"
    assert audit.target_type == "knowledge_document"
    assert audit.target_id == document.id


def test_disabled_channel_user_cannot_resolve_or_search():
    db = make_session()
    actor = Actor(channel="openclaw-weixin", external_user_id="disable-me@im.wechat")
    user = resolve_actor(db, actor)
    assert user is not None
    disable_user(db, user_id=user.id, actor_user_id="admin-user")

    assert resolve_actor(db, actor) is None


def test_search_knowledge_uses_hybrid_scores_for_related_questions():
    db = make_session()
    ingest_document(
        db,
        payload={
            "title": "售后发票规则",
            "content_text": "发票抬头错误时，客服需核对订单号、正确抬头和税号。未开票可直接修改，已开票需作废后重开。",
        },
        user_id="agent-1",
    )
    ingest_document(
        db,
        payload={
            "title": "物流丢件处理",
            "content_text": "物流疑似丢件时，客服先记录快递单号并联系快递核实。",
        },
        user_id="agent-1",
    )

    matches = search_knowledge(db, query="开票信息错了怎么处理", user=None, limit=2)

    assert matches[0]["document"].title == "售后发票规则"
    assert matches[0]["score"] > 0
    assert matches[0]["keyword_score"] >= 0
    assert matches[0]["vector_score"] > 0
    assert matches[0]["retrieval_method"] == "hybrid"


def test_find_source_files_matches_filename_and_document_title():
    db = make_session()
    ingest_document(
        db,
        payload={
            "title": "退换货政策原文",
            "content_text": "七天无理由退货需要商品不影响二次销售。",
            "project_id": "project-demo",
            "source_file_path": "/root/.openclaw/media/inbound/售后政策.docx",
            "source_file_name": "售后政策.docx",
            "source_file_storage": "openclaw_media",
        },
        user_id="agent-1",
    )

    by_file = find_source_files(db, query="售后政策", project_id="project-demo")
    by_title = find_source_files(db, query="退换货", project_id="project-demo")

    assert by_file[0]["source_file_name"] == "售后政策.docx"
    assert by_file[0]["source_file_path"] == "/root/.openclaw/media/inbound/售后政策.docx"
    assert by_file[0]["document_title"] == "退换货政策原文"
    assert by_title[0]["source_file_storage"] == "openclaw_media"


def test_find_source_files_returns_absolute_path_for_legacy_relative_upload():
    db = make_session()
    ingest_document(
        db,
        payload={
            "title": "历史上传原文",
            "content_text": "历史上传文件使用旧的相对路径保存。",
            "source_file_path": "uploads/knowledge_sources/legacy.txt",
            "source_file_name": "legacy.txt",
            "source_file_storage": "uploaded_copy",
        },
        user_id="agent-1",
    )

    result = find_source_files(db, query="legacy")

    assert Path(result[0]["source_file_path"]).is_absolute()
    assert result[0]["source_file_path"].endswith("uploads/knowledge_sources/legacy.txt")


def test_unanswered_question_status_can_be_resolved_or_ignored():
    db = make_session()
    answer_with_knowledge(db, query="会员生日券能不能叠加满减", user=None, actor_user_id="agent-1")
    unanswered = list_unanswered_questions(db)[0]

    updated = update_unanswered_question_status(db, unanswered["id"], status="resolved")

    assert updated["status"] == "resolved"
    assert list_unanswered_questions(db, status="pending") == []
    assert list_unanswered_questions(db, status="resolved")[0]["question"] == "会员生日券能不能叠加满减"


def test_import_markdown_and_plain_text_documents():
    db = make_session()

    markdown_result = import_faq_text(
        db,
        text="# 退货政策\n\n客户收到商品七天内可以申请无理由退货，特殊商品除外。",
        user_id="agent-1",
        source_type="markdown",
    )
    text_result = import_faq_text(
        db,
        text="物流异常处理\n客户反馈物流停滞超过48小时，客服需要联系快递核实并同步客户。",
        user_id="agent-1",
        source_type="txt",
    )

    assert markdown_result["imported_count"] == 1
    assert text_result["imported_count"] == 1
    titles = [document.title for document in db.query(KnowledgeDocument).order_by(KnowledgeDocument.title.asc()).all()]
    assert titles == ["物流异常处理", "退货政策"]


def test_import_knowledge_file_preserves_original_file_index(tmp_path):
    db = make_session()
    source = tmp_path / "退货说明.txt"
    source.write_text("退货说明\n客户收到商品七天内可以申请退货。", encoding="utf-8")

    result = import_knowledge_file(
        db,
        filename=source.name,
        content=source.read_bytes(),
        user_id="agent-1",
        source_type="txt",
    )

    assert result["imported_count"] == 1
    document = db.query(KnowledgeDocument).one()
    assert document.title == "退货说明"
    assert document.source_file_name == "退货说明.txt"
    assert document.source_file_path is not None
    assert Path(document.source_file_path).is_absolute()
    assert document.source_file_storage == "uploaded_copy"
    assert document.review_status == "pending_review"
    assert document.publication_status == "draft"
    assert document.index_status == "pending_review"
    assert "七天内可以申请退货" in document.content_text


def test_import_knowledge_file_reads_docx_and_xlsx(tmp_path):
    db = make_session()
    docx_path = tmp_path / "升级手册.docx"
    doc = Document()
    doc.add_paragraph("升级手册")
    doc.add_paragraph("升级前需要备份配置。")
    doc.save(docx_path)

    xlsx_path = tmp_path / "售后FAQ.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["问题", "答案"])
    sheet.append(["发票抬头写错", "未开票可修改，已开票需作废重开。"])
    workbook.save(xlsx_path)

    docx_result = import_knowledge_file(db, filename=docx_path.name, content=docx_path.read_bytes(), user_id="agent-1")
    xlsx_result = import_knowledge_file(db, filename=xlsx_path.name, content=xlsx_path.read_bytes(), user_id="agent-1")

    assert docx_result["imported_count"] == 1
    assert xlsx_result["imported_count"] == 1
    documents = db.query(KnowledgeDocument).order_by(KnowledgeDocument.title.asc()).all()
    assert {document.source_file_name for document in documents} == {"升级手册.docx", "售后FAQ.xlsx"}
    assert any("升级前需要备份配置" in document.content_text for document in documents)
    assert any("发票抬头写错" in document.content_text for document in documents)


def test_import_image_file_preserves_original_and_uses_ocr(monkeypatch):
    from app.services import support_operations_service

    db = make_session()
    monkeypatch.setattr(support_operations_service, "_extract_image_ocr_text", lambda content: "图片文字：退货请在七天内申请")

    result = import_knowledge_file(db, filename="退货规则.png", content=b"fake-image-bytes", user_id="agent-1")

    assert result["imported_count"] == 1
    document = db.query(KnowledgeDocument).one()
    assert document.title == "退货规则"
    assert document.source_file_name == "退货规则.png"
    assert document.source_file_mime_type == "image/png"
    assert document.source_file_storage == "uploaded_copy"
    assert "图片OCR识别文字" in document.content_text
    assert "退货请在七天内申请" in document.content_text


def test_upload_root_is_absolute_project_storage_path():
    from app.services import support_operations_service

    assert support_operations_service.UPLOAD_ROOT.is_absolute()
    assert support_operations_service.UPLOAD_ROOT.parts[-2:] == ("uploads", "knowledge_sources")


def test_import_knowledge_file_preserves_file_and_failure_record_when_parse_fails(tmp_path, monkeypatch):
    from app.services import support_operations_service

    db = make_session()
    monkeypatch.setattr(support_operations_service, "UPLOAD_ROOT", tmp_path / "knowledge_sources")

    result = import_knowledge_file(db, filename="空白.pdf", content=b"not-a-real-pdf", user_id="agent-1")

    document = db.query(KnowledgeDocument).one()
    assert result["ok"] is False
    assert document.parse_status == "parse_failed"
    assert document.publication_status == "draft"
    assert list((tmp_path / "knowledge_sources").glob("*"))


def test_teacher_tools_reuse_knowledge_search_permissions_and_export_docx(tmp_path, monkeypatch):
    from app.services import teacher_tools_service

    db = make_session()
    monkeypatch.setattr(teacher_tools_service, "TEACHER_EXPORT_ROOT", tmp_path / "teacher_exports")
    teacher = User(id="teacher-owner", name="教师", role="member", status="active", knowledge_access_policy="own")
    outsider = User(id="teacher-outsider", name="其他教师", role="member", status="active", knowledge_access_policy="own")
    db.add_all([teacher, outsider])
    db.commit()
    ingest_document(
        db,
        payload={
            "title": "函数讲义",
            "content_text": "函数定义域需要排除分母为0、偶次根号下小于0等情况。",
            "source_type": "teacher_material",
        },
        user_id="teacher-owner",
    )

    denied = search_teacher_materials(db, query="函数定义域", user=outsider)
    allowed = search_teacher_materials(db, query="函数定义域", user=teacher)
    questions = generate_teacher_questions(db, payload={"topic": "函数定义域", "count": 2}, user=teacher)
    paper = generate_teacher_paper(db, payload={"topic": "函数定义域", "question_counts": {"填空题": 1}}, user=teacher)
    handout = export_teacher_knowledge(db, payload={"topic": "函数定义域"}, user=teacher)

    assert denied["materials"] == []
    assert allowed["materials"][0]["document_title"] == "函数讲义"
    assert len(questions["questions"]) == 2
    assert paper["download"]["format"] == "docx"
    assert Path(paper["download"]["path"]).is_file()
    assert Path(handout["download"]["path"]).is_file()
    assert handout["handout"]["points"]


def test_import_docx_records_image_context_and_ocr(monkeypatch, tmp_path):
    from app.services import support_operations_service

    db = make_session()
    monkeypatch.setattr(support_operations_service, "_extract_image_ocr_text", lambda content: "截图文字：订单状态已发货")
    docx_path = tmp_path / "售后截图说明.docx"
    doc = Document()
    doc.add_paragraph("售后截图说明")
    doc.add_paragraph("下面的截图用于判断订单状态。")
    image_stream = BytesIO(base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="))
    doc.add_picture(image_stream)
    doc.add_paragraph("看到已发货后，需要引导客户关注物流。")
    doc.save(docx_path)

    result = import_knowledge_file(db, filename=docx_path.name, content=docx_path.read_bytes(), user_id="agent-1")

    assert result["imported_count"] == 1
    document = db.query(KnowledgeDocument).one()
    assert "图片上下文：下面的截图用于判断订单状态。 / 看到已发货后，需要引导客户关注物流。" in document.content_text
    assert "截图文字：订单状态已发货" in document.content_text
